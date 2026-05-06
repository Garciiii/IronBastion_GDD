# gerar_gdd.py
# Gerador do Game Design Document — Iron Bastion
# UC Gamificação | ISTEC Lisboa | 2024/25
# Autor: Grupo Iron Bastion
# Data: Abril 2025

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── CORES ────────────────────────────────────────────────────────────────────
AZUL_ESCURO   = RGBColor(0x1F, 0x38, 0x64)
AZUL_MEDIO    = RGBColor(0x2F, 0x54, 0x96)
AZUL_CLARO    = RGBColor(0x9D, 0xC3, 0xE6)
CINZENTO_CLARO= RGBColor(0xF2, 0xF2, 0xF2)
BRANCO        = RGBColor(0xFF, 0xFF, 0xFF)
PRETO         = RGBColor(0x00, 0x00, 0x00)
DOURADO       = RGBColor(0xC9, 0xA8, 0x4C)

# ─── HELPERS ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        tag = OxmlElement(f'w:{edge}')
        for key, val in attrs.items():
            tag.set(qn(f'w:{key}'), val)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, color=None, size=10, font='Calibri', align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def add_table(doc, headers, rows, col_widths=None):
    """Tabela com cabeçalho azul escuro e linhas alternadas."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabeçalho
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, AZUL_ESCURO)
        cell_text(c, h, bold=True, color=BRANCO, size=9)

    # Dados
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = CINZENTO_CLARO if ri % 2 == 0 else BRANCO
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg)
            cell_text(c, str(val), size=9)

    # Larguras
    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])
    return t

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = AZUL_ESCURO
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = AZUL_MEDIO
        run.font.size = Pt(13)
    elif level == 3:
        run.font.color.rgb = AZUL_MEDIO
        run.font.size = Pt(11)
    return p

def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return p

def add_mono(doc, text, size=8):
    """Bloco monospace para wireframes e ASCII."""
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(size)
    # Fundo cinzento claro via sombreado de parágrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return p

def add_numbered(doc, text, size=10):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return p

def add_page_break(doc):
    doc.add_page_break()

# ─── CABEÇALHO / RODAPÉ ───────────────────────────────────────────────────────

def setup_header_footer(doc):
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # Cabeçalho
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.text = ''
    run = hp.add_run('Iron Bastion — Game Design Document  |  UC Gamificação  |  ISTEC Lisboa 2024/25')
    run.font.size = Pt(8)
    run.font.color.rgb = AZUL_ESCURO
    run.font.name = 'Calibri'
    run.italic = True

    # Linha horizontal no cabeçalho
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Rodapé com número de página
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ''
    run = fp.add_run()
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    # Campo de número de página
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# ─── PÁGINA DE ROSTO ──────────────────────────────────────────────────────────

def add_cover(doc):
    # Espaço inicial
    for _ in range(4):
        doc.add_paragraph()

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('IRON BASTION')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = AZUL_ESCURO

    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Game Design Document — Versão 1.0')
    run.font.name = 'Calibri Light'
    run.font.size = Pt(18)
    run.font.color.rgb = AZUL_MEDIO

    doc.add_paragraph()

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"A última linha de defesa nunca dorme."')
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = DOURADO

    for _ in range(3):
        doc.add_paragraph()

    # Linha divisória
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)

    for _ in range(2):
        doc.add_paragraph()

    infos = [
        ('UC:', 'Gamificação'),
        ('Instituição:', 'ISTEC Lisboa'),
        ('Curso:', 'Licenciatura em Engenharia Informática'),
        ('Ano Letivo:', '2024/2025'),
        ('Data:', 'Abril 2025'),
    ]
    for label, val in infos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{label} ')
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r1.font.color.rgb = AZUL_ESCURO
        r2 = p.add_run(val)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'

    add_page_break(doc)

# ─── SECÇÕES ──────────────────────────────────────────────────────────────────

def sec_visao_geral(doc):
    add_heading(doc, '1. Visão Geral do Projeto', 1)
    add_body(doc, 'A tabela seguinte resume os principais atributos do jogo Iron Bastion.')
    doc.add_paragraph()
    add_table(doc,
        ['Atributo', 'Valor'],
        [
            ('Nome do jogo', 'Iron Bastion'),
            ('Tagline', '"A última linha de defesa nunca dorme."'),
            ('Género', 'Tower Defense 3D top-down'),
            ('Motor', 'Unity 2022.3 LTS'),
            ('Linguagem', 'C#'),
            ('Plataforma', 'PC — Windows (standalone build)'),
            ('Público-alvo', '13–35 anos, PEGI 12'),
            ('Classificação etária', 'PEGI 12 (violência fantástica, sem gore)'),
            ('Número de níveis', '3'),
            ('Número de torres', '4 (Arqueira, Canhão, Gelo, Arcana)'),
            ('Número de inimigos', '3 tipos (Comum, Veloz, Tanque)'),
            ('Duração média por sessão', '10–20 minutos por nível'),
            ('Modalidade', 'Trabalho de grupo'),
            ('Instituição', 'ISTEC Lisboa'),
        ],
        col_widths=[5, 11]
    )


def sec_narrativa(doc):
    add_heading(doc, '2. Narrativa e Universo', 1)

    add_heading(doc, '2.1 Mundo', 2)
    add_body(doc, 'Iron Bastion passa-se num mundo de fantasia industrial — um reino onde a magia coexiste com tecnologia rudimentar de vapor. Engrenagens e feitiços coexistem; canhões e varinhas mágicas são usados lado a lado. A estética 3D reflete esta dualidade: estruturas metálicas imponentes partilham espaço com runas mágicas incandescentes.')

    add_heading(doc, '2.2 Conflito', 2)
    add_body(doc, 'O Artefacto das Origens — um dispositivo mágico ancestral — foi ativado acidentalmente por um grupo de exploradores. O artefacto dispersou uma praga mágica que transforma humanos em Corrompidos: criaturas que perderam a razão e atacam tudo o que encontram. A praga alastra rapidamente, e o Bastião de Ferro é o último entreposto defensivo da civilização.')

    add_heading(doc, '2.3 Papel do Jogador', 2)
    add_body(doc, 'O jogador é o Comandante do Bastião de Ferro. A sua missão é coordenar a defesa do bastião, colocando e gerindo torres estrategicamente para impedir que os Corrompidos entrem e destruam o que resta da civilização.')

    add_heading(doc, '2.4 Os Corrompidos', 2)
    add_body(doc, 'Outrora humanos, agora monstruosidades movidas por instinto puro. Existem três variantes conforme o grau de corrupção:')
    add_bullet(doc, 'Corrompido Comum: totalmente transformado, andamento errático, pele esverdeada. Ameaça base mas numericamente significativo.')
    add_bullet(doc, 'Corrompido Veloz: corrupção acelerou o metabolismo, corpo magro e rápido. Difícil de atingir com torres de área.')
    add_bullet(doc, 'Corrompido Tanque: corrupção endureceu os tecidos, corpo massivo e blindado. Requer esforço concentrado de múltiplas torres.')

    add_heading(doc, '2.5 Apresentação da Narrativa', 2)
    add_bullet(doc, 'Texto de introdução antes do Nível 1 (narrativa de contexto)')
    add_bullet(doc, 'Descrições com lore nos tooltips de torres e inimigos')
    add_bullet(doc, 'Nomes temáticos por nível: O Portão do Norte / As Muralhas Caídas / O Coração do Bastião')
    add_bullet(doc, 'Textos de loading entre níveis com fragmentos de lore')


def sec_loop(doc):
    add_heading(doc, '3. Loop de Jogo Principal', 1)
    add_body(doc, 'O ciclo de jogo de Iron Bastion alterna entre fases de preparação estratégica e fases de combate ativo. O diagrama seguinte ilustra o fluxo completo:')
    doc.add_paragraph()
    add_mono(doc,
"""[Início do Nível]
      |
      v
[Fase de Preparação]
  --> Jogador coloca/melhora torres com ouro disponível
  --> Botão "Iniciar Vaga" disponível
      |
      v
[Iniciar Vaga] <-- Clique no botão OU tecla Space
      |
      v
[Vaga Ativa]
  --> Inimigos aparecem no ponto de spawn
  --> Seguem os waypoints do caminho (espaço 3D, Y=0)
  --> Torres disparam automaticamente nos inimigos dentro do alcance
      |
      +-- [Inimigo eliminado] ------> +Ouro
      +-- [Inimigo chega à base] ---> -Vidas
      |
      v
[Fim da Vaga]
  --> Bónus vaga perfeita se 0 inimigos chegaram (+20% ouro da vaga)
  --> Countdown de 15 seg para próxima vaga
      |
      v
[Mais vagas?]
  --> SIM: voltar à Fase de Preparação
  --> NAO: VITÓRIA --> calcular estrelas e pontuação
      |
[Vidas = 0?] --> GAME OVER em qualquer momento""")


def sec_torres(doc):
    add_heading(doc, '4. Sistema de Torres', 1)

    add_heading(doc, '4.1 Tabela Comparativa', 2)
    add_table(doc,
        ['Torre', 'Custo', 'Upgrade', 'Dano', 'Alcance', 'Cadência', 'Tipo', 'Especial'],
        [
            ('Arqueira',  '75',  '38',  '15 -> 22', '3.0u', '1.5 -> 2.1/s', 'Físico',     'Nv2: dispara 2 alvos'),
            ('Canhão',   '150', '75',  '80 -> 110', '2.5u', '0.4/s',         'Splash',    'Nv2: splash maior'),
            ('Gelo',     '125', '63',  '0',          '2.5u', '—',            'Slow',       'Slow 50% -> 70%'),
            ('Arcana',   '200', '100', '35 -> 43',  '3.5u', '0.8/s',         'Mágico+DoT', 'Amplifica vizinhas'),
        ],
        col_widths=[2.5, 1.5, 1.8, 2.5, 1.8, 2.8, 2.2, 3.7]
    )

    add_heading(doc, '4.2 Descrições Detalhadas', 2)

    torres = [
        ('Torre Arqueira (TowerArcher)', [
            'Papel: torre base, eficiente em custo, ideal para grupos de inimigos fracos.',
            'Nível 1: dispara 1 projétil rápido por alvo; dano 15; cadência 1.5/s.',
            'Nível 2: dispara simultaneamente para 2 alvos; dano 22; cadência 2.1/s.',
            'Sinergia recebida: +20% alcance se Torre Arcana adjacente.',
            'Visual placeholder 3D: cubo dourado/amarelo com ícone "A".',
            'Targeting: First (mais avançado no caminho).',
        ]),
        ('Torre Canhão (TowerCannon)', [
            'Papel: dano pesado em área; essencial contra grupos de Tanques.',
            'Nível 1: projétil lento com splash radius 1.2u; dano 80 em área (OverlapSphere 3D).',
            'Nível 2: splash radius aumenta para 1.8u; dano 110.',
            'Sinergia com Gelo: inimigos abrandados recebem +40% dano do Canhão.',
            'Visual placeholder 3D: cubo cinzento escuro.',
            'Targeting: First.',
        ]),
        ('Torre de Gelo (TowerIce)', [
            'Papel: suporte — abrandar inimigos para outras torres causarem mais dano.',
            'Não dispara projéteis — aplica slow contínuo por Physics.OverlapSphere a cada 0.5s.',
            'Nível 1: slow 50% de velocidade, duração 0.7s por pulso.',
            'Nível 2: slow 70% de velocidade, duração 0.7s, +0.5u alcance.',
            'Sinergia com Arcana adjacente: slow aumenta de 50% para 65%.',
            'Visual placeholder 3D: cubo azul claro / ciano.',
            'Não tem sistema de targeting (efeito de área passivo).',
        ]),
        ('Torre Arcana (TowerArcane)', [
            'Papel: dano médio com DoT + amplificadora de torres vizinhas.',
            'Nível 1: dano 35 + DoT 5 dano/seg durante 3s.',
            'Nível 2: dano 43 + DoT 8 dano/seg durante 4s.',
            'Sinergia fornecida: torres num raio de 2.5u ganham +20% alcance.',
            'Sinergia com Gelo: Torres de Gelo vizinhas passam de 50% → 65% slow.',
            'Visual placeholder 3D: cubo roxo / magenta.',
            'Targeting: Strongest (mais vida atual).',
        ]),
    ]
    for nome, pontos in torres:
        add_heading(doc, nome, 3)
        for pt in pontos:
            add_bullet(doc, pt)

    add_heading(doc, '4.3 Sistema de Sinergias', 2)
    add_body(doc, 'As sinergias incentivam o posicionamento estratégico das torres e são o diferencial central de Iron Bastion.')
    doc.add_paragraph()
    add_table(doc,
        ['Combinação', 'Bónus'],
        [
            ('Gelo + Canhão (adjacentes)', 'Canhão causa +40% dano a inimigos abrandados'),
            ('Arcana + qualquer (raio 2.5u)', 'Torres vizinhas ganham +20% alcance'),
            ('Arcana + Gelo (raio 2.5u)', 'Gelo passa de 50% -> 65% de slow'),
            ('Arqueira + Arqueira (adjacentes)', 'Cada uma ganha +10% cadência de fogo'),
        ],
        col_widths=[7, 9]
    )
    doc.add_paragraph()
    add_body(doc, 'Comunicação ao jogador:')
    add_bullet(doc, 'Linha visual 3D entre torres sinergéticas quando uma é selecionada (LineRenderer).')
    add_bullet(doc, 'Ícone de sinergia no painel de upgrade quando sinergia ativa.')
    add_bullet(doc, 'Tooltip explica o bónus ativo em linguagem clara.')

    add_heading(doc, '4.4 Sistema de Targeting', 2)
    add_body(doc, 'O jogador pode alterar o critério de targeting de cada torre individualmente:')
    add_table(doc,
        ['Modo', 'Comportamento'],
        [
            ('First', 'Ataca o inimigo mais avançado no caminho (padrão)'),
            ('Last', 'Ataca o inimigo mais atrasado no caminho'),
            ('Strongest', 'Ataca o inimigo com mais HP atual'),
            ('Weakest', 'Ataca o inimigo com menos HP atual (para terminar rapidamente)'),
        ],
        col_widths=[3.5, 12.5]
    )


def sec_inimigos(doc):
    add_heading(doc, '5. Sistema de Inimigos', 1)

    add_heading(doc, '5.1 Tabela Comparativa', 2)
    add_table(doc,
        ['Inimigo', 'HP', 'Velocidade', 'Recompensa', 'Dano à base', 'Res. Slow', 'Res. DoT'],
        [
            ('Corrompido Comum', '100', '2.0 u/s', '10', '1 vida', '0%', 'Nenhuma'),
            ('Corrompido Veloz', '50', '4.0 u/s', '6', '1 vida', '50%', 'Nenhuma'),
            ('Corrompido Tanque', '500', '0.8 u/s', '40', '3 vidas', '0%', 'Imune 2s após dano'),
        ],
        col_widths=[3.5, 1.5, 2.5, 2.8, 2.5, 2.0, 3.2]
    )

    add_heading(doc, '5.2 Fichas Detalhadas', 2)

    inimigos = [
        ('Corrompido Comum (EnemyCommon)', [
            'HP: 100 | Velocidade: 2.0 u/s | Recompensa: 10 ouro | Dano à base: 1 vida.',
            'Resistência a slow: 0% (slow total). Resistência a DoT: nenhuma.',
            'Comportamento: segue o caminho sem desvios; targeting básico.',
            'Visual placeholder 3D: cápsula verde (#4CAF50), escala 1.0.',
            'Barra de vida: visível após primeiro dano.',
        ]),
        ('Corrompido Veloz (EnemyFast)', [
            'HP: 50 | Velocidade: 4.0 u/s | Recompensa: 6 ouro | Dano à base: 1 vida.',
            'Resistência a slow: 50% (slow efetivo = metade do normal).',
            'Comportamento: corre; chega antes dos Comuns; difícil de atingir com Canhão.',
            'Visual placeholder 3D: cápsula amarela (#FFC107), escala 0.8x (mais pequena e rápida).',
        ]),
        ('Corrompido Tanque (EnemyTank)', [
            'HP: 500 | Velocidade: 0.8 u/s | Recompensa: 40 ouro | Dano à base: 3 vidas.',
            'Resistência a slow: 0%. Resistência a DoT: imune durante 2s após tomar qualquer dano.',
            'Comportamento: avança devagar; absorve enorme quantidade de dano; prioridade máxima.',
            'Visual placeholder 3D: cápsula vermelha escura (#B71C1C), escala 1.5x (maior e intimidante).',
            'Barra de vida: sempre visível (pela sua importância estratégica).',
        ]),
    ]
    for nome, pontos in inimigos:
        add_heading(doc, nome, 3)
        for pt in pontos:
            add_bullet(doc, pt)


def sec_vagas(doc):
    add_heading(doc, '6. Sistema de Vagas', 1)

    add_heading(doc, '6.1 Estrutura de uma Vaga', 2)
    add_table(doc,
        ['Parâmetro', 'Valor'],
        [
            ('Intervalo entre spawns', '0.8 segundos'),
            ('Tempo de preparação entre vagas', '15 segundos (countdown visível)'),
            ('Bónus de vaga perfeita', '+20% do ouro total gerado nessa vaga'),
        ],
        col_widths=[7, 9]
    )

    add_heading(doc, '6.2 Nível 1 — O Portão do Norte (5 vagas)', 2)
    add_table(doc,
        ['Vaga', 'Comuns', 'Velozes', 'Tanques', 'Ouro base'],
        [
            ('1', '8', '0', '0', '80'),
            ('2', '10', '2', '0', '112'),
            ('3', '12', '4', '0', '144'),
            ('4', '8', '6', '1', '208'),
            ('5 (boss)', '5', '5', '1 *', '250'),
        ],
        col_widths=[2.5, 2.5, 2.5, 2.5, 2.5]
    )
    add_body(doc, '* Vaga boss: Tanque com HP = 750 (1.5x normal).')

    add_heading(doc, '6.3 Nível 2 — As Muralhas Caídas (8 vagas)', 2)
    add_table(doc,
        ['Vaga', 'Comuns', 'Velozes', 'Tanques', 'Ouro base'],
        [
            ('1', '10', '3', '0', '118'),
            ('2', '12', '5', '0', '150'),
            ('3', '10', '6', '1', '210'),
            ('4', '8', '8', '2', '268'),
            ('5', '15', '5', '1', '240'),
            ('6', '10', '10', '2', '300'),
            ('7', '8', '8', '3', '344'),
            ('8 (boss)', '5', '5', '2 *', '350'),
        ],
        col_widths=[2.5, 2.5, 2.5, 2.5, 2.5]
    )
    add_body(doc, '* Vaga boss: Tanques com HP = 900.')

    add_heading(doc, '6.4 Nível 3 — O Coração do Bastião (12 vagas)', 2)
    add_table(doc,
        ['Vaga', 'Comuns', 'Velozes', 'Tanques', 'Ouro base'],
        [
            ('1', '12', '5', '0', '150'),
            ('2', '15', '8', '1', '250'),
            ('3', '10', '10', '2', '300'),
            ('4', '8', '12', '3', '362'),
            ('5', '20', '5', '2', '330'),
            ('6', '12', '12', '3', '396'),
            ('7', '10', '15', '4', '450'),
            ('8', '15', '10', '4', '420'),
            ('9', '8', '20', '3', '448'),
            ('10', '20', '10', '5', '550'),
            ('11', '15', '15', '5', '600'),
            ('12 (boss)', '10', '10', '3 *', '700'),
        ],
        col_widths=[2.5, 2.5, 2.5, 2.5, 2.5]
    )
    add_body(doc, '* Vaga boss final: Tanques com HP = 1500 (3x normal).')


def sec_economia(doc):
    add_heading(doc, '7. Sistema de Economia', 1)
    add_table(doc,
        ['Parâmetro', 'Valor'],
        [
            ('Ouro inicial — Nível 1', '150'),
            ('Ouro inicial — Nível 2', '200'),
            ('Ouro inicial — Nível 3', '250'),
            ('Venda de torre', '50% do total investido (custo + upgrades)'),
            ('Bónus vaga perfeita', '+20% do ouro gerado na vaga'),
            ('Transferência entre níveis', 'Não — cada nível começa do zero'),
        ],
        col_widths=[7, 9]
    )


def sec_vidas_pontuacao(doc):
    add_heading(doc, '8. Sistema de Vidas e Pontuação', 1)

    add_heading(doc, '8.1 Vidas', 2)
    add_body(doc, 'Cada nível começa com 10 vidas. Não existe recuperação de vidas durante o nível.')
    add_table(doc,
        ['Inimigo', 'Vidas perdidas'],
        [
            ('Corrompido Comum', '-1 vida'),
            ('Corrompido Veloz', '-1 vida'),
            ('Corrompido Tanque', '-3 vidas'),
        ],
        col_widths=[8, 8]
    )

    add_heading(doc, '8.2 Sistema de Estrelas', 2)
    add_table(doc,
        ['Classificação', 'Condição'],
        [
            ('3 Estrelas', '8 a 10 vidas restantes no fim do nível'),
            ('2 Estrelas', '4 a 7 vidas restantes'),
            ('1 Estrela',  '1 a 3 vidas restantes'),
        ],
        col_widths=[5, 11]
    )

    add_heading(doc, '8.3 Pontuação', 2)
    add_table(doc,
        ['Evento', 'Pontos'],
        [
            ('Base por nível completo', '+1000 pts'),
            ('Bónus por estrela extra (acima de 1)', '+500 pts'),
            ('Bónus por vaga perfeita', '+100 pts por vaga'),
            ('Pontuação máxima teórica (Nível 1)', '~2600 pts'),
        ],
        col_widths=[9, 7]
    )
    add_body(doc, 'O high score é guardado com PlayerPrefs e persiste entre sessões.')


def sec_regras(doc):
    add_heading(doc, '9. Regras do Jogo', 1)
    regras = [
        'O jogador vence se sobreviver a todas as vagas com pelo menos 1 vida.',
        'Game Over se as vidas chegarem a 0.',
        'Torres só podem ser colocadas em células "Buildable" (não no caminho dos inimigos).',
        'Cada célula suporta apenas 1 torre.',
        'Torres NÃO podem ser colocadas durante vagas ativas.',
        'Torres PODEM ser vendidas a qualquer momento, inclusive durante vagas.',
        'Ouro não transfere entre níveis — cada nível começa do zero.',
        'Upgrades são permanentes e irreversíveis.',
        'Vender uma torre devolve 50% do total investido (custo base + custo de upgrades).',
        'Bónus de vaga perfeita: concedido apenas se 0 inimigos chegaram à base nessa vaga.',
    ]
    for r in regras:
        add_numbered(doc, r)


def sec_niveis(doc):
    add_heading(doc, '10. Design dos Níveis', 1)
    add_body(doc, 'Os mapas são grelhas 3D com câmara perspetiva top-down. Os inimigos movem-se ao nível Y=0, e as torres são colocadas nas células Buildable. Os layouts abaixo representam uma vista de cima.')

    add_heading(doc, '10.1 Nível 1 — O Portão do Norte', 2)
    add_table(doc,
        ['Parâmetro', 'Valor'],
        [
            ('Tema', 'Entrada norte da fortaleza, campo aberto, caminho simples'),
            ('Torres disponíveis', 'Arqueira, Canhão'),
            ('Vagas', '5'),
            ('Inimigos', 'Comum, Veloz'),
            ('Ouro inicial', '150'),
            ('Grid', '15x10 células (1 célula = 1 unidade Unity)'),
        ],
        col_widths=[5, 11]
    )
    doc.add_paragraph()
    add_mono(doc,
"""  0 1 2 3 4 5 6 7 8 9 0 1 2 3 4
0 . . . . . . . . . . . . . . .
1 . . . . # # # # # # # # . . .
2 . . . . # . . . . . . . . . .
3 . . . . # . . . . . . . . . .
4 . . . . # # # # # # # # . . .
5 . . . . . . . . . . . . # . .
6 . . . . . . . . . . . . # . .
7 . . . . . . . . . . . . # . .
8 . . . . . . . . . . . . # . .
9 S . . . . . . . . . . . E . .

# = caminho dos inimigos
S = Ponto de Spawn  |  E = Ponto de Entrada (base)
. = célula construível""")
    doc.add_paragraph()
    add_body(doc, 'Waypoints (em ordem, coordenadas world):')
    add_bullet(doc, 'WP0 Spawn: (-1, 0, 9)  |  WP1: (4, 0, 9)  |  WP2: (4, 0, 4)  |  WP3: (12, 0, 4)  |  WP4 End: (12, 0, 9)')

    add_heading(doc, '10.2 Nível 2 — As Muralhas Caídas', 2)
    add_table(doc,
        ['Parâmetro', 'Valor'],
        [
            ('Tema', 'Muralhas em ruínas, caminho em S com 2 curvas'),
            ('Torres disponíveis', 'Arqueira, Canhão, Gelo'),
            ('Vagas', '8'),
            ('Inimigos', 'Comum, Veloz, Tanque'),
            ('Ouro inicial', '200'),
            ('Grid', '20x12 células'),
        ],
        col_widths=[5, 11]
    )
    doc.add_paragraph()
    add_mono(doc,
"""  0 1 2 3 4 5 6 7 8 9 0 1 2 3 4 5 6 7 8 9
0 . . . . . . . . . . . . . . . . . . . .
1 S # # # # # . . . . . . . . . . . . . .
2 . . . . . # . . . . . . . . . . . . . .
3 . . . . . # # # # # # . . . . . . . . .
4 . . . . . . . . . . # . . . . . . . . .
5 . . . . . . . . . . # # # # # # . . . .
6 . . . . . . . . . . . . . . . # . . . .
7 . . . . . . . . . . . . . . . # . . . .
8 . . . . . . . . . . . . . . . # . . . .
9 . . . . . . . . . . . . . . . # . . . .
0 . . . . . . . . . . . . . . . # . . . .
1 . . . . . . . . . . . . . . . E . . . .""")
    doc.add_paragraph()
    add_body(doc, 'Waypoints: WP0(-1,0,1) WP1(5,0,1) WP2(5,0,3) WP3(10,0,3) WP4(10,0,5) WP5(15,0,5) WP6(15,0,11) WP7(15,0,12)')

    add_heading(doc, '10.3 Nível 3 — O Coração do Bastião', 2)
    add_table(doc,
        ['Parâmetro', 'Valor'],
        [
            ('Tema', 'Centro da fortaleza, caminho labiríntico com 4 curvas'),
            ('Torres disponíveis', 'Todas (Arqueira, Canhão, Gelo, Arcana)'),
            ('Vagas', '12'),
            ('Inimigos', 'Todos os tipos (inclui vaga boss final)'),
            ('Ouro inicial', '250'),
            ('Grid', '22x14 células — caminho em espiral'),
        ],
        col_widths=[5, 11]
    )
    doc.add_paragraph()
    add_mono(doc,
"""  0 1 2 3 4 5 6 7 8 9 0 1 2 3 4 5 6 7 8 9 0 1
0 . . . . . . . . . . . . . . . . . . . . . .
1 S # # # # # # # # # # # # # # # # . . . . .
2 . . . . . . . . . . . . . . . . # . . . . .
3 . . . . . . . . . . . . . . . . # . . . . .
4 . . . # # # # # # # # # # # # # . . . . . .
5 . . . # . . . . . . . . . . . . . . . . . .
6 . . . # . . . . . . . . . . . . . . . . . .
7 . . . # # # # # # # # # . . . . . . . . . .
8 . . . . . . . . . . . # . . . . . . . . . .
9 . . . . . . . . . . . # # # # # . . . . . .
0 . . . . . . . . . . . . . . . # . . . . . .
1 . . . . . . . . . . . . . . . # . . . . . .
2 . . . . . . . . . . . . . . . # . . . . . .
3 . . . . . . . . . . . . . . . E . . . . . .""")
    doc.add_paragraph()
    add_body(doc, 'Waypoints: WP0(-1,0,1) WP1(16,0,1) WP2(16,0,4) WP3(3,0,4) WP4(3,0,7) WP5(11,0,7) WP6(11,0,9) WP7(15,0,9) WP8(15,0,13) WP9(15,0,14)')


def sec_ui(doc):
    add_heading(doc, '11. Interface do Utilizador', 1)
    add_body(doc, 'Todos os ecrãs seguem a paleta de cores definida (fundo #0D1B2A, acentos dourados #C9A84C). Os wireframes abaixo representam a estrutura funcional.')

    add_heading(doc, '11.1 HUD em Jogo', 2)
    add_mono(doc,
"""+-----------------------------------------------------------------+
|  [vida] 10    [ouro] 150        VAGA 3 / 5       [PAUSA]       |
+-----------------------------------------------------------------+
|                                                                 |
|                      [AREA DE JOGO 3D]                         |
|              (camera perspetiva top-down)                       |
|                                                                 |
+----------------------------------+--------------------------+---+
|  [PAINEL DE TORRES]              | [PAINEL DE UPGRADE]         |
|  Arqueira   75   [COMPRAR]       | (aparece ao clicar numa     |
|  Canhao    150   [COMPRAR]       |  torre existente)           |
|  Gelo      125   [BLOQUEADO]     |                             |
|  Arcana    200   [BLOQUEADO]     |  Torre: Arqueira Nv.1       |
|                                  |  Dano: 15 -> 22             |
|  [INICIAR VAGA]                  |  Custo upgrade: 38          |
|                                  |  [MELHORAR]  [VENDER 38]   |
+----------------------------------+-----------------------------+""")

    add_heading(doc, '11.2 Menu Principal', 2)
    add_mono(doc,
"""+-----------------------------------------------------------------+
|                                                                 |
|                      IRON BASTION                               |
|           "A ultima linha de defesa nunca dorme."               |
|                                                                 |
|                     [ JOGAR ]                                   |
|                     [ CREDITOS ]                                |
|                     [ SAIR ]                                    |
|                                                                 |
|                                          v1.0 - ISTEC 2024/25  |
+-----------------------------------------------------------------+""")

    add_heading(doc, '11.3 Seleção de Nível', 2)
    add_mono(doc,
"""+-----------------------------------------------------------------+
|  <- VOLTAR               SELECIONAR NIVEL                       |
+-----------------------------------------------------------------+
|  +---------------+  +---------------+  +---------------+       |
|  |   NIVEL 1     |  |   NIVEL 2     |  |   NIVEL 3     |       |
|  | O Portao do   |  | As Muralhas   |  | O Coracao do  |       |
|  |   Norte       |  |   Caidas      |  |   Bastiao     |       |
|  |  * * (vazio)  |  |  * (vazio x2) |  |  [BLOQUEADO]  |       |
|  |   [JOGAR]     |  |   [JOGAR]     |  |               |       |
|  +---------------+  +---------------+  +---------------+       |
+-----------------------------------------------------------------+""")

    add_heading(doc, '11.4 Menu de Pausa', 2)
    add_mono(doc,
"""+-----------------------------------------------------------------+
|                          PAUSA                                  |
|                                                                 |
|                     [ CONTINUAR ]                               |
|                     [ RECOMECAR NIVEL ]                         |
|                     [ MENU PRINCIPAL ]                          |
|                                                                 |
+-----------------------------------------------------------------+""")

    add_heading(doc, '11.5 Ecrã de Vitória', 2)
    add_mono(doc,
"""+-----------------------------------------------------------------+
|                      BASTIAO DEFENDIDO!                         |
|                                                                 |
|                        * * (vazio)                              |
|                                                                 |
|    Pontuacao: 1850 pts    Vidas: 6/10    Vagas perfeitas: 3/5  |
|                                                                 |
|            [ PROXIMO NIVEL ]         [ MENU ]                   |
+-----------------------------------------------------------------+""")

    add_heading(doc, '11.6 Ecrã de Derrota', 2)
    add_mono(doc,
"""+-----------------------------------------------------------------+
|                      O BASTIAO CAIU...                          |
|                                                                 |
|                  Resististe ate a Vaga 3 de 5                   |
|                                                                 |
|            [ TENTAR NOVAMENTE ]         [ MENU ]                |
+-----------------------------------------------------------------+""")


def sec_controlos(doc):
    add_heading(doc, '12. Controlos', 1)
    add_table(doc,
        ['Ação', 'Controlo'],
        [
            ('Selecionar torre no shop', 'Clique esquerdo no botão'),
            ('Colocar torre no mapa', 'Clique esquerdo em célula válida'),
            ('Cancelar colocação', 'Clique direito OU Escape'),
            ('Selecionar torre existente', 'Clique esquerdo na torre'),
            ('Desselecionar', 'Clique direito OU clicar fora'),
            ('Iniciar vaga', 'Botão "Iniciar Vaga" OU barra de espaço'),
            ('Pausar / Retomar', 'Botão pausa OU Escape (durante jogo)'),
            ('Upgrade de torre', 'Botão "Melhorar" no painel'),
            ('Vender torre', 'Botão "Vender" no painel'),
        ],
        col_widths=[7, 9]
    )


def sec_arte(doc):
    add_heading(doc, '13. Arte e Estética', 1)

    add_heading(doc, '13.1 Estilo Visual', 2)
    add_body(doc, 'Fantasia industrial — steampunk sombrio com elementos mágicos. O jogo é 3D com câmara perspetiva top-down fixa. As torres e inimigos são representados por modelos 3D simples (cubos/cápsulas) em fase de protótipo, prontos para substituição por modelos finais.')

    add_heading(doc, '13.2 Paleta de Cores', 2)
    add_table(doc,
        ['Elemento', 'Cor', 'Hex'],
        [
            ('Fundo de UI', 'Azul muito escuro', '#0D1B2A'),
            ('Acentos UI', 'Dourado', '#C9A84C'),
            ('Texto principal', 'Branco acinzentado', '#E8E8E8'),
            ('Torres — metálico', 'Cinzento azulado', '#607D8B'),
            ('Torres — dourado', 'Amarelo quente', '#FFC107'),
            ('Inimigos — corrupção', 'Verde doentio', '#4CAF50'),
            ('Inimigos — magia', 'Roxo intenso', '#7B1FA2'),
            ('Caminho', 'Castanho terra', '#795548'),
            ('Área construível', 'Verde muito escuro', '#1B5E20'),
            ('Highlight válido', 'Verde translúcido', '#4CAF50 70%'),
            ('Highlight inválido', 'Vermelho translúcido', '#F44336 70%'),
        ],
        col_widths=[5, 5, 6]
    )

    add_heading(doc, '13.3 Assets de Placeholder 3D', 2)
    add_body(doc, 'Todos os modelos podem ser gerados programaticamente em Unity sem dependências externas:')
    add_bullet(doc, 'Torres: GameObject com MeshRenderer (cubo) + material colorido + texto 3D (letra A/C/G/X)')
    add_bullet(doc, 'Inimigos: GameObject com MeshRenderer (cápsula) + material colorido')
    add_bullet(doc, 'Caminho: planos 3D com material castanho')
    add_bullet(doc, 'Projéteis: esfera 3D pequena com material emissivo')

    add_heading(doc, '13.4 Assets Externos Recomendados (CC0)', 2)
    add_bullet(doc, 'Kenney.nl — Tower Defense Top-Down Pack (sprites adaptáveis a texturas)')
    add_bullet(doc, 'Kenney.nl — UI Pack (botões, painéis)')
    add_bullet(doc, 'OpenGameArt.org — músicas de fantasia')
    add_bullet(doc, 'Unity Asset Store — Synty Studios packs (opcionais, baixo custo)')


def sec_arquitetura(doc):
    add_heading(doc, '14. Arquitetura Técnica', 1)

    add_heading(doc, '14.1 Padrões de Design', 2)
    add_table(doc,
        ['Padrão', 'Uso em Iron Bastion'],
        [
            ('Singleton', 'GameManager — acesso global ao estado do jogo'),
            ('Observer / Events', 'GameEvents — comunicação desacoplada entre sistemas'),
            ('Object Pool', 'Inimigos e projéteis — evita Instantiate/Destroy em Update'),
            ('ScriptableObject', 'Dados de torres, inimigos e vagas — configuráveis no Inspector'),
            ('Strategy', 'Targeting das torres — First/Last/Strongest/Weakest'),
            ('Inheritance', 'EnemyBase → tipos; TowerBase → tipos'),
        ],
        col_widths=[4, 12]
    )

    add_heading(doc, '14.2 Considerações 3D vs 2D', 2)
    add_table(doc,
        ['Sistema', 'Implementação 3D'],
        [
            ('Física', 'Physics.OverlapSphere (em vez de Physics2D.OverlapCircle)'),
            ('Câmara', 'Perspetiva top-down (em vez de Ortográfica 2D)'),
            ('Coordenadas', 'Vector3 com Y=0 para movimentos no plano'),
            ('Renderização', 'MeshRenderer + Material (em vez de SpriteRenderer)'),
            ('Waypoints', 'Transform.position no espaço 3D'),
            ('Raycasting', 'Physics.Raycast para seleção de células no mapa'),
        ],
        col_widths=[4, 12]
    )

    add_heading(doc, '14.3 Estrutura de Pastas', 2)
    add_mono(doc,
"""Assets/
├── Scripts/
│   ├── Core/
│   │   ├── GameManager.cs         <- Singleton; estado global
│   │   ├── WaveManager.cs         <- Gestao de vagas e spawning
│   │   └── GameEvents.cs          <- Sistema de eventos global
│   ├── Enemies/
│   │   ├── EnemyBase.cs           <- Classe base; movimento 3D; dano; morte
│   │   ├── EnemyCommon.cs         <- Corrompido Comum
│   │   ├── EnemyFast.cs           <- Corrompido Veloz
│   │   ├── EnemyTank.cs           <- Corrompido Tanque
│   │   └── EnemySpawner.cs        <- Instanciar inimigos no spawn point
│   ├── Towers/
│   │   ├── TowerBase.cs           <- Classe base; targeting; disparo; upgrade
│   │   ├── TowerArcher.cs         <- Arqueira; multi-shot nivel 2
│   │   ├── TowerCannon.cs         <- Canhao; splash damage 3D
│   │   ├── TowerIce.cs            <- Gelo; slow por area (OverlapSphere); sem projeteis
│   │   └── TowerArcane.cs         <- Arcana; DoT; amplificacao de vizinhas
│   ├── Projectiles/
│   │   ├── ProjectileBase.cs      <- Projetil base; homing 3D; dano single
│   │   └── ProjectileSplash.cs    <- Projetil com dano em area (Canhao)
│   ├── Grid/
│   │   ├── GridManager.cs         <- Grid de celulas 3D; colocacao; highlight
│   │   └── GridCell.cs            <- Estado de cada celula do grid
│   ├── Path/
│   │   └── PathManager.cs         <- Lista de waypoints 3D; Gizmos
│   ├── UI/
│   │   ├── UIManager.cs           <- HUD; paineis; eventos de UI
│   │   ├── TowerShopUI.cs         <- Botoes de compra; modo de colocacao
│   │   ├── TowerInfoPanel.cs      <- Painel de upgrade/venda
│   │   └── HealthBar.cs           <- Barra de vida sobre inimigos (Canvas World Space)
│   └── Utils/
│       └── ObjectPool.cs          <- Pool generico para inimigos e projeteis
├── ScriptableObjects/
│   ├── TowerData.cs
│   ├── EnemyData.cs
│   └── WaveData.cs
├── Data/
│   ├── Towers/  (TowerData_Archer/Cannon/Ice/Arcane.asset)
│   ├── Enemies/ (EnemyData_Common/Fast/Tank.asset)
│   └── Waves/   (Level1-3, W1-W12)
├── Prefabs/
│   ├── Enemies/  ├── Towers/  ├── Projectiles/  └── UI/
├── Scenes/
│   ├── MainMenu.unity  ├── Level1.unity  ├── Level2.unity  └── Level3.unity
└── Audio/
    ├── Music/  └── SFX/""")

    add_heading(doc, '14.4 Tags e Layers', 2)
    add_mono(doc,
"""Tags:   Enemy | Tower | Projectile | Waypoint
Layers: Enemy (6) | Tower (7) | Buildable (8)""")

    add_heading(doc, '14.5 Configuração de Câmara 3D', 2)
    add_table(doc,
        ['Parâmetro', 'Valor'],
        [
            ('Tipo', 'Perspetiva (Field of View: 60°)'),
            ('Rotação', '-60° no eixo X (top-down ligeiramente inclinado)'),
            ('Posição Nível 1', 'Centrada no mapa, Y=8'),
            ('Posição Nível 2', 'Centrada no mapa, Y=10'),
            ('Posição Nível 3', 'Centrada no mapa, Y=12'),
            ('Background', 'Cor sólida #0D1B2A (azul muito escuro)'),
        ],
        col_widths=[5, 11]
    )

    add_heading(doc, '14.6 Build Settings', 2)
    add_mono(doc,
"""Index 0: MainMenu
Index 1: Level1
Index 2: Level2
Index 3: Level3""")


def sec_glossario(doc):
    add_heading(doc, '15. Glossário', 1)
    termos = [
        ('Tower Defense', 'Género de jogo onde o jogador constrói torres para impedir inimigos de atravessar um caminho.'),
        ('Waypoint', 'Ponto de passagem que define o caminho dos inimigos no mapa 3D.'),
        ('Object Pool', 'Padrão de otimização que reutiliza objetos já criados em vez de instanciar e destruir constantemente.'),
        ('ScriptableObject', 'Asset de dados Unity que permite configurar parâmetros de jogo sem modificar código.'),
        ('Singleton', 'Padrão de design que garante a existência de uma única instância de uma classe.'),
        ('Observer / Events', 'Padrão que permite comunicação desacoplada entre sistemas através de eventos.'),
        ('DoT', 'Damage over Time — dano aplicado gradualmente ao longo do tempo.'),
        ('Splash', 'Dano em área — afeta todos os inimigos dentro de um raio.'),
        ('Slow', 'Efeito que reduz a velocidade de movimento dos inimigos.'),
        ('HUD', 'Heads-Up Display — interface sobreposta ao jogo com informação essencial.'),
        ('Vaga', 'Wave — grupo de inimigos que ataca em conjunto.'),
        ('Spawn', 'Ponto/momento em que os inimigos aparecem no mapa.'),
        ('Targeting', 'Sistema que define qual inimigo uma torre ataca prioritariamente.'),
        ('Sinergia', 'Bónus obtido pela combinação estratégica de torres próximas.'),
        ('Corrompidos', 'Inimigos do jogo — antigos humanos transformados pela praga mágica.'),
        ('Physics.OverlapSphere', 'Função Unity 3D que deteta todos os colliders dentro de uma esfera no espaço 3D.'),
        ('MeshRenderer', 'Componente Unity que renderiza um modelo 3D com um material atribuído.'),
    ]
    add_table(doc,
        ['Termo', 'Definição'],
        termos,
        col_widths=[4, 12]
    )


def sec_estado(doc):
    add_heading(doc, '16. Estado Atual e Próximos Passos', 1)

    add_heading(doc, '16.1 Estado das Entregas', 2)
    add_table(doc,
        ['Entrega', 'Estado', 'Descrição'],
        [
            ('Entrega 1', 'Concluída', 'Proposta do jogo em documento Word'),
            ('Entrega 2', 'Em curso', 'GDD completo (este documento)'),
            ('Entrega 3', 'A fazer', 'Protótipo inicial jogável — Nível 1, Torres Arqueira + Canhão'),
            ('Entrega 4', 'A fazer', 'Versão final — 3 níveis, GDD final, Build.zip + Project.zip + PDF'),
        ],
        col_widths=[3, 3, 10]
    )

    add_heading(doc, '16.2 Próximos Passos', 2)
    add_numbered(doc, 'Configurar projeto Unity 2022.3 LTS como projeto 3D.')
    add_numbered(doc, 'Implementar scripts C# (ordem: GameEvents → GameManager → Path/Grid → Enemies → Towers → UI).')
    add_numbered(doc, 'Criar ScriptableObjects e configurar assets de dados.')
    add_numbered(doc, 'Construir o Nível 1 no Unity com waypoints e grid 3D.')
    add_numbered(doc, 'Testar Entrega 3 e gerar build executável.')
    add_numbered(doc, 'Implementar Níveis 2 e 3 com todas as torres e inimigos.')
    add_numbered(doc, 'Polir UI, adicionar áudio e preparar Entrega 4.')

    add_heading(doc, '16.3 Referências', 2)
    add_heading(doc, 'Jogos de Referência', 3)
    refs_jogos = [
        'Kingdom Rush (2011, Ironhide) — UI de Tower Defense, sistema de estrelas.',
        'Bloons TD 6 (2018, Ninja Kiwi) — upgrades em árvore, variedade de inimigos.',
        'Plants vs. Zombies (2009, PopCap) — acessibilidade, curva de aprendizagem.',
        'Dungeon Warfare 2 (2017, Valsar) — sinergias entre torres, posicionamento.',
        'Defenders Quest (2012, Level Up Labs) — progressão RPG em Tower Defense.',
    ]
    for r in refs_jogos:
        add_bullet(doc, r)

    add_heading(doc, 'Referências Técnicas e Académicas', 3)
    refs_tec = [
        'Unity Documentation (2022.3 LTS) — https://docs.unity3d.com',
        'Schell, J. (2019). The Art of Game Design: A Book of Lenses. CRC Press.',
        'Fullerton, T. (2018). Game Design Workshop. CRC Press.',
        'Deterding et al. (2011). From game design elements to gamefulness. MindTrek.',
        'Kenney.nl — Assets CC0 para protótipo.',
    ]
    for r in refs_tec:
        add_bullet(doc, r)


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    setup_header_footer(doc)

    # Estilos base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_cover(doc)

    sec_visao_geral(doc)
    doc.add_paragraph()
    sec_narrativa(doc)
    add_page_break(doc)

    sec_loop(doc)
    doc.add_paragraph()
    sec_torres(doc)
    add_page_break(doc)

    sec_inimigos(doc)
    doc.add_paragraph()
    sec_vagas(doc)
    add_page_break(doc)

    sec_economia(doc)
    doc.add_paragraph()
    sec_vidas_pontuacao(doc)
    doc.add_paragraph()
    sec_regras(doc)
    add_page_break(doc)

    sec_niveis(doc)
    add_page_break(doc)

    sec_ui(doc)
    add_page_break(doc)

    sec_controlos(doc)
    doc.add_paragraph()
    sec_arte(doc)
    add_page_break(doc)

    sec_arquitetura(doc)
    add_page_break(doc)

    sec_glossario(doc)
    doc.add_paragraph()
    sec_estado(doc)

    output = r'C:/Users/helde/IronBastion_GDD/Entrega2_GDD_IronBastion.docx'
    doc.save(output)
    print(f'GDD gerado com sucesso: {output}')

    import os
    size = os.path.getsize(output)
    print(f'Tamanho: {size / 1024:.1f} KB')


if __name__ == '__main__':
    main()
